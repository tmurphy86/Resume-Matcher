"""Tests for the ATS-safe .docx export service (RH-208).

All unit tests are purely deterministic: build_docx is a pure function that
takes a ResumeData and returns bytes — no LLM, no network, no mocks needed.
The integration test uses the isolated_db fixture for a real but disposable DB.
"""

import io
import json

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient

from app.schemas.models import ResumeData
from app.services.docx_export import build_docx


# ---------------------------------------------------------------------------
# Test data helpers
# ---------------------------------------------------------------------------

def _minimal_resume_data(**overrides: object) -> ResumeData:
    """Build a minimal ResumeData for testing."""
    base: dict = {
        "personalInfo": {
            "name": "Jane Doe",
            "title": "Software Engineer",
            "email": "jane@example.com",
            "phone": "+1-555-0100",
            "location": "San Francisco, CA",
            "website": "https://janedoe.dev",
            "linkedin": "linkedin.com/in/janedoe",
            "github": "github.com/janedoe",
        },
        "summary": "Experienced software engineer with a passion for clean code.",
        "workExperience": [
            {
                "id": 1,
                "title": "Senior Engineer",
                "company": "Acme Corp",
                "location": "San Francisco, CA",
                "years": "Jan 2021 - Present",
                "description": [
                    "Built scalable REST APIs with Python and FastAPI",
                    "Led migration to microservices",
                ],
            }
        ],
        "education": [
            {
                "id": 1,
                "institution": "MIT",
                "degree": "B.S. Computer Science",
                "years": "2014 - 2018",
                "description": "Graduated with honors",
            }
        ],
        "personalProjects": [
            {
                "id": 1,
                "name": "OpenAPI Generator",
                "role": "Creator",
                "years": "2021 - Present",
                "github": "github.com/janedoe/openapi-gen",
                "description": [
                    "CLI tool generating API clients",
                    "500+ GitHub stars",
                ],
            }
        ],
        "additional": {
            "technicalSkills": ["Python", "FastAPI", "Docker"],
            "languages": ["English (Native)"],
            "certificationsTraining": ["AWS Solutions Architect"],
            "awards": ["Employee of the Year 2022"],
        },
        "customSections": {},
        "sectionMeta": [],
    }
    base.update(overrides)
    return ResumeData.model_validate(base)


def _get_all_text(doc: Document) -> str:
    """Return all paragraph text joined by newlines."""
    return "\n".join(p.text for p in doc.paragraphs)


def _get_paragraph_styles(doc: Document) -> list[str]:
    """Return the style name of each paragraph."""
    return [p.style.name for p in doc.paragraphs]


# ---------------------------------------------------------------------------
# Unit tests — pure function, no mocks
# ---------------------------------------------------------------------------

def test_build_docx_returns_valid_docx() -> None:
    """build_docx must return valid OOXML bytes that python-docx can open."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    assert isinstance(raw, bytes)
    assert len(raw) > 0
    doc = Document(io.BytesIO(raw))
    assert len(doc.paragraphs) > 0


def test_build_docx_contains_name() -> None:
    """The candidate's name must appear in the document."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)
    assert "Jane Doe" in full_text


def test_build_docx_contains_contact_info() -> None:
    """Email and phone must appear in the contact paragraph."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)
    assert "jane@example.com" in full_text
    assert "+1-555-0100" in full_text


def test_build_docx_section_order() -> None:
    """Summary, Experience, Education, Skills, Projects must appear in order."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)

    summary_pos = full_text.index("Summary")
    experience_pos = full_text.index("Experience")
    education_pos = full_text.index("Education")
    skills_pos = full_text.index("Skills")
    projects_pos = full_text.index("Projects")

    assert summary_pos < experience_pos < education_pos < skills_pos < projects_pos


def test_build_docx_uses_heading_styles() -> None:
    """Section headers must use real Word heading styles, not plain paragraphs."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    styles = _get_paragraph_styles(doc)
    # Must have at least Heading 1 and Heading 2 present
    assert "Heading 1" in styles, "Missing Heading 1 style"
    assert "Heading 2" in styles, "Missing Heading 2 style"


def test_build_docx_uses_list_bullet_style() -> None:
    """Experience bullet points must use 'List Bullet' style, not manual prefix."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    styles = _get_paragraph_styles(doc)
    assert "List Bullet" in styles, "Missing List Bullet style — bullets must use real list style"


def test_build_docx_no_manual_bullet_prefix() -> None:
    """No paragraph text should start with a manual '• ' prefix."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    for para in doc.paragraphs:
        text = para.text.strip()
        assert not text.startswith("•"), (
            f"Found manual bullet prefix in paragraph: {text!r}"
        )
        assert not text.startswith("- ") or para.style.name == "List Bullet", (
            f"Potential manual dash-bullet in non-list paragraph: {text!r}"
        )


def test_build_docx_bullets_from_description() -> None:
    """Legacy resume with description[] must produce List Bullet paragraphs."""
    data = _minimal_resume_data(
        workExperience=[
            {
                "id": 1,
                "title": "Engineer",
                "company": "TechCo",
                "years": "2020 - 2023",
                "description": ["Shipped feature A", "Reduced latency by 30%"],
            }
        ]
    )
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)
    assert "Shipped feature A" in full_text
    assert "Reduced latency by 30%" in full_text

    bullet_texts = [
        p.text for p in doc.paragraphs if p.style.name == "List Bullet"
    ]
    assert "Shipped feature A" in bullet_texts
    assert "Reduced latency by 30%" in bullet_texts


def test_build_docx_experience_heading_contains_company() -> None:
    """The experience Heading 2 must include the company name."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    heading2_texts = [
        p.text for p in doc.paragraphs if p.style.name == "Heading 2"
    ]
    # At least one Heading 2 must include "Acme Corp"
    assert any("Acme Corp" in t for t in heading2_texts), (
        f"Heading 2 paragraphs: {heading2_texts}"
    )


def test_build_docx_education_heading_contains_institution() -> None:
    """The education Heading 2 must include the institution name."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    heading2_texts = [
        p.text for p in doc.paragraphs if p.style.name == "Heading 2"
    ]
    assert any("MIT" in t for t in heading2_texts), (
        f"Heading 2 paragraphs: {heading2_texts}"
    )


def test_build_docx_skills_section() -> None:
    """Technical skills must appear in the Skills section."""
    data = _minimal_resume_data()
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)
    assert "Python" in full_text
    assert "FastAPI" in full_text
    assert "Docker" in full_text


def test_build_docx_empty_experience_is_omitted() -> None:
    """A resume with no work experience must not include the Experience heading."""
    data = _minimal_resume_data(workExperience=[])
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    # Check Heading 1 paragraphs only — the word "Experience" may appear in
    # other text (e.g. summary "Experienced...") so we target the heading style.
    heading1_texts = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert "Experience" not in heading1_texts, (
        f"'Experience' Heading 1 should be absent when workExperience is empty; "
        f"got headings: {heading1_texts}"
    )


def test_build_docx_empty_projects_is_omitted() -> None:
    """A resume with no projects must not include the Projects heading."""
    data = _minimal_resume_data(personalProjects=[])
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)
    assert "Projects" not in full_text


def test_build_docx_summary_omitted_when_empty() -> None:
    """A resume with an empty summary must not include the Summary heading."""
    data = _minimal_resume_data(summary="")
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)
    assert "Summary" not in full_text


def test_build_docx_minimal_resume_no_crash() -> None:
    """A resume with only personalInfo and empty sections must not crash."""
    data = ResumeData.model_validate(
        {
            "personalInfo": {"name": "Minimal Person", "email": "min@example.com"},
            "summary": "",
            "workExperience": [],
            "education": [],
            "personalProjects": [],
            "additional": {},
            "customSections": {},
            "sectionMeta": [],
        }
    )
    raw = build_docx(data)
    doc = Document(io.BytesIO(raw))
    full_text = _get_all_text(doc)
    assert "Minimal Person" in full_text


# ---------------------------------------------------------------------------
# Integration test — real router via httpx ASGI transport + isolated_db
# ---------------------------------------------------------------------------

@pytest.mark.integration
async def test_endpoint_exports_docx(isolated_db: object) -> None:
    """GET /{resume_id}/export/docx must return 200 with correct content-type.

    Flow: seed the isolated DB with a resume that has processed_data, then
    hit the endpoint and validate the response bytes are a valid .docx.
    """
    from app.main import app

    processed_data = _minimal_resume_data().model_dump()

    # Seed a resume directly into the isolated DB
    resume = await isolated_db.create_resume(  # type: ignore[union-attr]
        content=json.dumps(processed_data),
        content_type="json",
        filename="test_resume.json",
        is_master=True,
        processed_data=processed_data,
        processing_status="ready",
    )
    resume_id = resume["resume_id"]

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        response = await client.get(f"/api/v1/resumes/{resume_id}/export/docx")

    assert response.status_code == 200
    content_type = response.headers.get("content-type", "")
    assert "wordprocessingml" in content_type, f"Unexpected content-type: {content_type}"

    # Verify the returned bytes are a valid .docx
    doc = Document(io.BytesIO(response.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane Doe" in full_text


@pytest.mark.integration
async def test_endpoint_404_for_missing_resume(isolated_db: object) -> None:
    """GET /{resume_id}/export/docx must return 404 when resume does not exist."""
    from app.main import app

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        response = await client.get("/api/v1/resumes/nonexistent-id/export/docx")

    assert response.status_code == 404


@pytest.mark.integration
async def test_endpoint_422_for_resume_without_processed_data(
    isolated_db: object,
) -> None:
    """GET /{resume_id}/export/docx must return 422 when resume has no processed_data."""
    from app.main import app

    # Seed a resume without processed_data (e.g. failed processing)
    resume = await isolated_db.create_resume(  # type: ignore[union-attr]
        content="# Some markdown content",
        content_type="md",
        filename="raw_resume.pdf",
        is_master=True,
        processed_data=None,
        processing_status="failed",
    )
    resume_id = resume["resume_id"]

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        response = await client.get(f"/api/v1/resumes/{resume_id}/export/docx")

    assert response.status_code == 422
