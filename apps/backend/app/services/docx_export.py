"""ATS-safe .docx export from ResumeData via python-docx (ADR-004).

Goal: structural fidelity for ATS parsers and human editors, NOT pixel fidelity.
Real Word heading styles and real bullet lists — no manual "• " prefix strings,
no text boxes, no floating elements.
"""

import io
import logging

from docx import Document
from docx.shared import Pt

from app.schemas.models import AdditionalInfo, Education, Experience, PersonalInfo, Project, ResumeData

logger = logging.getLogger(__name__)


def build_docx(resume_data: ResumeData) -> bytes:
    """Build an ATS-safe .docx from ResumeData. Returns raw bytes.

    Uses python-docx with real Word heading styles and proper list paragraphs.
    The generated document is intentionally plain — ATS parsers and human
    editors are the target audience, not visual renderers.
    """
    doc = Document()
    _add_contact_header(doc, resume_data.personalInfo)
    if resume_data.summary:
        _add_summary(doc, resume_data.summary)
    if resume_data.workExperience:
        _add_experience(doc, resume_data.workExperience)
    if resume_data.education:
        _add_education(doc, resume_data.education)
    if resume_data.additional:
        _add_skills(doc, resume_data.additional)
    if resume_data.personalProjects:
        _add_projects(doc, resume_data.personalProjects)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_contact_header(doc: Document, info: PersonalInfo) -> None:
    """Add name as bold paragraph and contact line as plain text.

    ADR-004: avoid the Murphy template's black header bar / gray band.
    Approximate simply with bold name paragraph, plain contact text.
    """
    if info.name:
        name_para = doc.add_paragraph()
        run = name_para.add_run(info.name)
        run.bold = True
        run.font.size = Pt(16)

    contact_parts: list[str] = []
    if info.title:
        contact_parts.append(info.title)
    if info.email:
        contact_parts.append(info.email)
    if info.phone:
        contact_parts.append(info.phone)
    if info.location:
        contact_parts.append(info.location)
    if info.website:
        contact_parts.append(info.website)
    if info.linkedin:
        contact_parts.append(info.linkedin)
    if info.github:
        contact_parts.append(info.github)

    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))


def _add_summary(doc: Document, summary: str) -> None:
    """Add Summary section with Heading 2 and paragraph text."""
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary)


def _add_experience(doc: Document, experiences: list[Experience]) -> None:
    """Add Experience section with Heading 1; each job uses Heading 2 + bullets.

    Bullet text source: Experience.description is the resolved list of bullet
    strings. When bullet_blocks is non-empty the ResumeData model_validator
    rebuilds description from the active variant's text; otherwise the legacy
    description list is used as-is. Either way, this function only reads
    experience.description — the schema layer handles block resolution.
    """
    doc.add_heading("Experience", level=1)
    for exp in experiences:
        # Company + dates as Heading 2
        company_line = exp.company
        if exp.years:
            company_line = f"{company_line}  |  {exp.years}" if company_line else exp.years
        if company_line:
            doc.add_heading(company_line, level=2)

        # Role as italic paragraph
        if exp.title:
            role_para = doc.add_paragraph()
            run = role_para.add_run(exp.title)
            run.italic = True

        # Location as plain paragraph
        if exp.location:
            doc.add_paragraph(exp.location)

        # Bullet points — real List Bullet style (ATS-safe, no manual "• ")
        for bullet in exp.description:
            if bullet.strip():
                doc.add_paragraph(bullet.strip(), style="List Bullet")


def _add_education(doc: Document, education: list[Education]) -> None:
    """Add Education section with Heading 1; each entry as Heading 2 + text."""
    doc.add_heading("Education", level=1)
    for edu in education:
        # Institution + years as Heading 2
        institution_line = edu.institution
        if edu.years:
            institution_line = (
                f"{institution_line}  |  {edu.years}" if institution_line else edu.years
            )
        if institution_line:
            doc.add_heading(institution_line, level=2)

        # Degree as italic paragraph
        if edu.degree:
            degree_para = doc.add_paragraph()
            run = degree_para.add_run(edu.degree)
            run.italic = True

        # Description as plain paragraph
        if edu.description:
            doc.add_paragraph(edu.description)


def _add_skills(doc: Document, additional: AdditionalInfo) -> None:
    """Add Skills & Awards section with Heading 1.

    Each category (technical skills, languages, etc.) is rendered as a
    comma-separated paragraph under a bold label — ATS-safe and compact.
    """
    doc.add_heading("Skills & Awards", level=1)

    if additional.technicalSkills:
        para = doc.add_paragraph()
        para.add_run("Technical Skills: ").bold = True
        para.add_run(", ".join(additional.technicalSkills))

    if additional.languages:
        para = doc.add_paragraph()
        para.add_run("Languages: ").bold = True
        para.add_run(", ".join(additional.languages))

    if additional.certificationsTraining:
        para = doc.add_paragraph()
        para.add_run("Certifications & Training: ").bold = True
        para.add_run(", ".join(additional.certificationsTraining))

    if additional.awards:
        para = doc.add_paragraph()
        para.add_run("Awards: ").bold = True
        para.add_run(", ".join(additional.awards))


def _add_projects(doc: Document, projects: list[Project]) -> None:
    """Add Projects section with Heading 1; each project as Heading 2 + bullets."""
    doc.add_heading("Projects", level=1)
    for proj in projects:
        # Name + years as Heading 2
        name_line = proj.name
        if proj.years:
            name_line = f"{name_line}  |  {proj.years}" if name_line else proj.years
        if name_line:
            doc.add_heading(name_line, level=2)

        # Role as italic paragraph
        if proj.role:
            role_para = doc.add_paragraph()
            run = role_para.add_run(proj.role)
            run.italic = True

        # Links as plain paragraph
        links: list[str] = []
        if proj.github:
            links.append(proj.github)
        if proj.website:
            links.append(proj.website)
        if links:
            doc.add_paragraph(" | ".join(links))

        # Bullet points — real List Bullet style
        for bullet in proj.description:
            if bullet.strip():
                doc.add_paragraph(bullet.strip(), style="List Bullet")
